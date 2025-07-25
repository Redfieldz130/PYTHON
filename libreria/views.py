from django.shortcuts import render, redirect, get_object_or_404
from .forms import EquipoForm, AsignacionForm, CustomUserCreationForm
from .models import Equipo, Asignacion, HistorialEquipo
from django.contrib import messages
from django.contrib.auth import logout, login, authenticate
from django.http import JsonResponse, HttpResponse
from datetime import datetime
from django.utils import timezone
import openpyxl
from docx import Document
from openpyxl.styles import Font
from django.db.models.functions import TruncDate
from django.db.models import DateField
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt


def historico_equipos(request):
    historial = HistorialEquipo.objects.all().order_by('-fecha')
    modelos = HistorialEquipo.objects.filter(modelo__isnull=False, modelo__gt='').values('modelo').distinct()
    tipos = HistorialEquipo.objects.filter(tipo__isnull=False).values('tipo').distinct()
    modelos_formateados = [modelo['modelo'] for modelo in modelos]
    tipos_formateados = [{'valor': tipo['tipo'], 'display': tipo['tipo']} for tipo in tipos]
    fechas_formateadas = [item.fecha.strftime('%Y-%m-%d') for item in historial if item.fecha]
    return render(request, 'Equipos/historico.html', {
        'historial': historial,
        'fechas_con_registros': fechas_formateadas,
        'modelos': modelos_formateados,
        'tipos': tipos_formateados
    })

def detalle_equipo_json(request, equipo_id):
    equipo = get_object_or_404(Equipo, pk=equipo_id)
    data = {
        'modelo': equipo.modelo or 'N/A',
        'marca': equipo.marca or 'N/A',
        'serial': equipo.serial or 'N/A',
        'mac_address': equipo.mac_address or 'N/A',
        'estado': equipo.get_estado_display() or 'N/A',  # Usar get_estado_display()
        'observaciones': equipo.observaciones or 'Sin observaciones',
        'tipo': equipo.get_tipo_display() or 'N/A',
        'fecha_fabricacion': equipo.fecha_fabricacion.strftime('%Y-%m-%d') if equipo.fecha_fabricacion else 'N/A',
        'vida_util_anios': str(equipo.vida_util_anios) if equipo.vida_util_anios else 'N/A',
        'nombre_red': equipo.nombre_red or 'N/A',
        'ubicacion': equipo.ubicacion or 'N/A',
        'empleado_responsable': equipo.empleado_responsable or 'N/A',
        'proveedor': equipo.proveedor or 'N/A',
        'valor_compra': str(equipo.valor_compra) if equipo.valor_compra else 'N/A',
        'fecha_compra': equipo.fecha_compra.strftime('%Y-%m-%d') if equipo.fecha_compra else 'N/A',
        'fecha_recepcion': equipo.fecha_recepcion.strftime('%Y-%m-%d') if equipo.fecha_recepcion else 'N/A',
        'fecha_mantenimiento': equipo.fecha_mantenimiento.strftime('%Y-%m-%d') if equipo.fecha_mantenimiento else 'N/A',
        'size_pantalla': str(equipo.size_pantalla) if equipo.size_pantalla else 'N/A',
        'resolucion': equipo.resolucion or 'N/A',
        'procesador_marca': equipo.procesador_marca or 'N/A',
        'procesador_velocidad': str(equipo.procesador_velocidad) if equipo.procesador_velocidad else 'N/A',
        'procesador_generacion': equipo.procesador_generacion or 'N/A',
        'sistema_operativo': equipo.sistema_operativo or 'N/A',
        'sistema_operativo_version': equipo.sistema_operativo_version or 'N/A',
        'sistema_operativo_bits': equipo.sistema_operativo_bits or 'N/A',
        'almacenamiento_capacidad': str(equipo.almacenamiento_capacidad) if equipo.almacenamiento_capacidad else 'N/A',
        'memoria': str(equipo.memoria) if equipo.memoria else 'N/A',
        'impresora_tipo': equipo.impresora_tipo or 'N/A',
        'impresora_velocidad_ppm': str(equipo.impresora_velocidad_ppm) if equipo.impresora_velocidad_ppm else 'N/A',
        'impresora_color': equipo.impresora_color or 'N/A',  # Usar directamente
        'impresora_conexion': equipo.impresora_conexion or 'N/A',
        'cpu_formato_diseno': equipo.cpu_formato_diseno or 'N/A',
        'proyector_lumens': str(equipo.proyector_lumens) if equipo.proyector_lumens else 'N/A',
        'ups_vatios': str(equipo.ups_vatios) if equipo.ups_vatios else 'N/A',
        'ups_fecha_bateria': equipo.ups_fecha_bateria.strftime('%Y-%m-%d') if equipo.ups_fecha_bateria else 'N/A',
        'scanner_velocidad': str(equipo.scanner_velocidad) if equipo.scanner_velocidad else 'N/A',
        'scanner_color': equipo.scanner_color or 'N/A',  # Usar directamente
        'pantalla_proyector_tipo': equipo.pantalla_proyector_tipo or 'N/A',
        'server_numero_procesadores': str(equipo.server_numero_procesadores) if equipo.server_numero_procesadores else 'N/A',
        'licencia_tipo': equipo.licencia_tipo or 'N/A',
        'licencia_clase': equipo.licencia_clase or 'N/A',
        'mouse_tipo': equipo.mouse_tipo or 'N/A',
        'mouse_conexion': equipo.mouse_conexion or 'N/A',
        'clase_disco': equipo.clase_disco or 'N/A',
    }
    return JsonResponse(data)

def eliminar_equipos_seleccionados(request):
    if request.method == 'POST':
        equipo_ids = request.POST.getlist('equipos_seleccionados')
        equipos = Equipo.objects.filter(id__in=equipo_ids)
        if not equipos.exists():
            messages.error(request, 'No se seleccionaron equipos para eliminar.')
            return redirect('equipos')
        equipos.delete()
        messages.success(request, 'Los equipos seleccionados fueron eliminados con éxito.')
        return redirect('equipos')
    return redirect('equipos')

def actualizar_estado(request, equipo_id):
    equipo = get_object_or_404(Equipo, id=equipo_id)
    equipo.estado = 'Asignado'
    equipo.save()
    return JsonResponse({'estado': equipo.estado})

def Inicio(request):
    return render(request, 'Paginas/inicio.html')

def nosotros(request):
    return render(request, 'Paginas/nosotros.html')

def listar_equipos(request):
    equipos_todos = Equipo.objects.all()
    equipos_categorias = {
        'equipos': ['laptop', 'impresora', 'cpu', 'monitor', 'proyector', 'ups', 'scanner', 'pantalla_proyector', 'tablet', 'server', 'router', 'access_point', 'camara_web', 'disco_duro'],
        'accesorios': ['mouse', 'teclado', 'headset', 'bocina', 'brazo_monitor', 'memoria_usb', 'pointer', 'kit_herramientas', 'generador_tono', 'tester', 'multimetro'],
        'licencias': ['licencia_informatica'],
        'materiales': ['cartucho', 'toner', 'botella_tinta'],
    }
    equipos_por_categoria = {
        'todos': equipos_todos,
        'equipos': equipos_todos.filter(tipo__in=equipos_categorias['equipos']),
        'accesorios': equipos_todos.filter(tipo__in=equipos_categorias['accesorios']),
        'licencias': equipos_todos.filter(tipo__in=equipos_categorias['licencias']),
        'materiales': equipos_todos.filter(tipo__in=equipos_categorias['materiales']),
    }
    asignaciones_dict = {asignacion.equipo.id: asignacion.id for asignacion in Asignacion.objects.filter(fecha_final__isnull=True)}
    print("Asignaciones_dict:", asignaciones_dict)  # Para depuración
    context = {
        'equipos_por_categoria': equipos_por_categoria,
        'asignaciones_dict': asignaciones_dict,
    }
    return render(request, 'Equipos/Index.html', context)
def crear_equipos(request):
    if request.method == 'POST':
        formulario = EquipoForm(request.POST)
        if formulario.is_valid():
            equipo = formulario.save()
            HistorialEquipo.objects.create(
                equipo=equipo,
                accion='creado',
                usuario=request.user,
                tipo=equipo.tipo,
                marca=equipo.marca,
                modelo=equipo.modelo,
                serial=equipo.serial,
            )
            messages.success(request, "Equipo agregado exitosamente.")
            return redirect('equipos')
        else:
            messages.error(request, "Error en el formulario. Por favor, revise los campos.")
    else:
        formulario = EquipoForm()
    return render(request, 'Equipos/Crear.html', {'formulario': formulario})

def editar_equipo(request, equipo_id):
    equipo = get_object_or_404(Equipo, id=equipo_id)
    if request.method == 'POST':
        formulario = EquipoForm(request.POST, instance=equipo)
        if formulario.is_valid():
            equipo = formulario.save()
            HistorialEquipo.objects.create(
                equipo=equipo,
                accion='editado',
                usuario=request.user,
                tipo=equipo.tipo,
                marca=equipo.marca,
                modelo=equipo.modelo,
                serial=equipo.serial,
            )
            messages.success(request, f'El equipo "{equipo.modelo}" ha sido actualizado con éxito.')
            return redirect('equipos')
        else:
            messages.error(request, "Error en el formulario. Por favor, revise los campos.")
    else:
        formulario = EquipoForm(instance=equipo)
    return render(request, 'Equipos/Editar.html', {'formulario': formulario, 'equipo': equipo})

def asignar(request):
    tipos_equipo = Equipo.objects.values_list('tipo', flat=True).distinct().order_by('tipo')
    equipos_asignados = Asignacion.objects.filter(fecha_final__isnull=True).values_list('equipo_id', flat=True)
    equipos = Equipo.objects.exclude(id__in=equipos_asignados).filter(estado='Disponible')
    asignaciones = Asignacion.objects.all()

    if request.method == 'POST':
        print(request.POST)  # Depuración: imprime los datos enviados
        formulario = AsignacionForm(request.POST)
        if formulario.is_valid():
            print(formulario.cleaned_data)  # Depuración: imprime los datos validados
            equipo = formulario.cleaned_data['equipo']
            if Asignacion.objects.filter(equipo=equipo, fecha_final__isnull=True).exists():
                messages.error(request, f'El equipo "{equipo.modelo}" ya está asignado actualmente.')
                return redirect('asignar')
            asignacion = formulario.save()
            print(f"Asignación creada: ID={asignacion.id}, Cédula={asignacion.colaborador_cedula}")  # Depuración
            equipo.estado = 'Asignado'
            equipo.save()
            HistorialEquipo.objects.create(
                equipo=equipo,
                accion='asignado',
                usuario=request.user,
                tipo=equipo.tipo,
                marca=equipo.marca,
                modelo=equipo.modelo,
                serial=equipo.serial,
            )
            messages.success(request, f'El equipo "{equipo.modelo}" ha sido asignado a {asignacion.colaborador_nombre} (Cédula: {asignacion.colaborador_cedula}).')
            return redirect('asignar')
        else:
            messages.error(request, 'Error en el formulario de asignación. Por favor, revise los campos.')
            print(formulario.errors)  # Depuración: imprime errores del formulario
    else:
        formulario = AsignacionForm()
    return render(request, 'Equipos/Asignar.html', {
        'equipos': equipos,
        'asignaciones': asignaciones,
        'tipos_equipo': tipos_equipo,
        'fecha_hoy': timezone.now().date(),
        'formulario': formulario,
    })
def Listadeasignados(request):
    asignaciones = Asignacion.objects.all()
    return render(request, 'Equipos/Listadeasignados.html', {'asignaciones': asignaciones})

def desasignar(request, id):
    print(f"Intentando desasignar asignación con ID: {id}")  # Depuración
    try:
        asignacion = Asignacion.objects.get(id=id)
        equipo = asignacion.equipo
        asignacion.delete()
        equipo.estado = "Disponible"
        equipo.save()
        HistorialEquipo.objects.create(
            equipo=equipo,
            accion='desasignado',
            usuario=request.user,
            tipo=equipo.tipo,
            marca=equipo.marca,
            modelo=equipo.modelo,
            serial=equipo.serial,
        )
        messages.success(request, f'El equipo "{equipo.modelo}" ha sido desasignado.')
    except Asignacion.DoesNotExist:
        print(f"No se encontró asignación con ID: {id}")  # Depuración
        messages.error(request, 'No se pudo encontrar la asignación.')
    return redirect('equipos')  # Cambiado de 'asignar' a 'equipos'

@require_POST
def borrar_equipo(request, equipo_id):
    equipo = get_object_or_404(Equipo, id=equipo_id)
    if request.method == 'POST':
        try:
            HistorialEquipo.objects.create(
                equipo=None,
                accion='eliminado',
                usuario=request.user,
                tipo=equipo.tipo,
                marca=equipo.marca,
                modelo=equipo.modelo,
                serial=equipo.serial,
            )
            equipo.delete()
            return JsonResponse({'status': 'success', 'message': 'Equipo eliminado correctamente'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': f'Error al eliminar el equipo: {str(e)}'}, status=500)
    return JsonResponse({'status': 'error', 'message': 'Método no permitido'}, status=405)

def liberar_equipo(request, equipo_id):
    equipo = get_object_or_404(Equipo, id=equipo_id)
    equipo.estado = "Disponible"
    equipo.save()
    messages.success(request, f'Equipo {equipo.modelo} ahora está disponible.')
    return redirect('equipos')

def registro(request):
    if request.method == 'POST':
        formulario = CustomUserCreationForm(request.POST)
        if formulario.is_valid():
            formulario.save()
            user = authenticate(username=formulario.cleaned_data["username"], password=formulario.cleaned_data["password1"])
            login(request, user)
            messages.success(request, "¡Registro exitoso!")
            return redirect('inicio')
        else:
            messages.error(request, "Error en el formulario de registro. Por favor, revise los campos.")
    else:
        formulario = CustomUserCreationForm()
    return render(request, 'registration/registro.html', {'form': formulario})


def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            messages.success(request, f'¡Bienvenido {user.first_name}!')
            return redirect('inicio')
        else:
            messages.error(request, 'Usuario o contraseña inválidos')
    return render(request, 'registration/login.html')

def logout_view(request):
    logout(request)
    messages.success(request, "Has cerrado sesión correctamente.")
    return redirect("login")

def inicio_view(request):
    return render(request, "inicio.html")

def generar_constancia(request, asignacion_id):
    asignacion = get_object_or_404(Asignacion, id=asignacion_id)
    nombre_empleado = asignacion.colaborador_nombre
    equipo = asignacion.equipo
    fecha_entrega = asignacion.fecha_entrega.strftime("%d de %B de %Y")
    fecha_final = asignacion.fecha_final.strftime("%d de %B de %Y") if asignacion.fecha_final else "No aplica"
    mac_address = equipo.mac_address if equipo.mac_address else "No disponible"
    colaborador_cedula = asignacion.colaborador_cedula if asignacion.colaborador_cedula else "No disponible"
    print(f"Colaborador Cedula: '{colaborador_cedula}'")  # Depuración
    plantilla_path = 'C:/Users/efrain.delacruz/Desktop/PYTHON-1/static/Word.docx'
    doc = Document(plantilla_path)
    reemplazar_texto(doc, '{{ nombreempleado }}', nombre_empleado)
    reemplazar_texto(doc, '{{ equipodescripcion }}', f'{equipo.marca} {equipo.modelo}')
    reemplazar_texto(doc, '{{ equiposerial }}', equipo.serial)
    reemplazar_texto(doc, '{{ macaddress }}', mac_address)
    reemplazar_texto(doc, '{{ fecha_entrega }}', fecha_entrega)
    reemplazar_texto(doc, '{{ fecha_final }}', fecha_final)
    reemplazar_texto(doc, '{{ cedulaempleado }}', colaborador_cedula)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=constancia_salida_{equipo.id}.docx'
    doc.save(response)
    return response

def reemplazar_texto(doc, marcador, texto):
    """ Reemplaza texto en el documento, incluyendo párrafos y tablas. """
    texto = str(texto)  # Asegura que el texto sea string
    for p in doc.paragraphs:
        if marcador in p.text:
            print(f"Reemplazando '{marcador}' con '{texto}' en párrafo: {p.text}")  # Depuración
            p.text = p.text.replace(marcador, texto)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if marcador in p.text:
                        print(f"Reemplazando '{marcador}' con '{texto}' en celda: {p.text}")  # Depuración
                        p.text = p.text.replace(marcador, texto)
def exportar_excel(request):
    categoria = request.GET.get('categoria', 'todos')
    equipos = Equipo.objects.all()
    if categoria != 'todos':
        categorias_dict = {
            'equipos': ['laptop', 'impresora', 'cpu', 'monitor', 'proyector', 'ups', 'scanner', 'pantalla_proyector', 'tablet', 'server', 'router', 'access_point', 'camara_web', 'disco_duro'],
            'accesorios': ['mouse', 'teclado', 'headset', 'bocina', 'brazo_monitor', 'memoria_usb', 'pointer', 'kit_herramientas', 'generador_tono', 'tester', 'multimetro'],
            'licencias': ['licencia_informatica'],
            'materiales': ['cartucho', 'toner', 'botella_tinta'],
        }
        if categoria in categorias_dict:
            equipos = equipos.filter(tipo__in=categorias_dict[categoria])
        else:
            equipos = Equipo.objects.none()
    # Crear un libro de trabajo y una hoja
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Inventario de Tecnología"

    # Definir encabezados
    headers = [
        'Tipo', 'Marca', 'Modelo', 'Serial', 'Estado', 'Ubicación',
        'Fecha de Compra', 'Valor de Compra', 'Empleado Responsable', 'Observaciones'
    ]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.font = Font(bold=True)

    # Llenar datos
    for row_num, equipo in enumerate(equipos, 2):
        ws.cell(row=row_num, column=1).value = equipo.get_tipo_display()
        ws.cell(row=row_num, column=2).value = equipo.marca
        ws.cell(row=row_num, column=3).value = equipo.modelo
        ws.cell(row=row_num, column=4).value = equipo.serial
        ws.cell(row=row_num, column=5).value = equipo.estado
        ws.cell(row=row_num, column=6).value = equipo.ubicacion
        ws.cell(row=row_num, column=7).value = equipo.fecha_compra.strftime('%Y-%m-%d') if equipo.fecha_compra else ''
        ws.cell(row=row_num, column=8).value = str(equipo.valor_compra) if equipo.valor_compra else ''
        ws.cell(row=row_num, column=9).value = equipo.empleado_responsable
        ws.cell(row=row_num, column=10).value = equipo.observaciones

    # Ajustar el ancho de las columnas
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column].width = adjusted_width

    # Preparar la respuesta HTTP
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="Inventario_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx"'
    
    # Guardar el libro en la respuesta
    wb.save(response)
    return response

# libreria/views.py
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Equipo, HistorialEquipo
from .forms import EquipoForm
import openpyxl
from datetime import datetime
import re

def importar_excel(request):
    if request.method == 'POST':
        archivo = request.FILES.get('archivo_excel')
        if not archivo:
            messages.error(request, 'No se ha seleccionado ningún archivo.')
            return redirect('equipos')

        try:
            wb = openpyxl.load_workbook(archivo)
            ws = wb.active
            errors = []
            success_count = 0

            # Mapear encabezados del Excel a campos del modelo
            header_map = {
                'tipo': ['tipo', 'type', 'category', 'tipo_categoria'],
                'marca': ['marca', 'brand'],
                'modelo': ['modelo', 'model'],
                'serial': ['serial', 'serie', 'numero_serie', 'numero de serie'],
                'estado': ['estado', 'status', 'estado_del_producto'],
                'ubicacion': ['ubicacion', 'ubicación', 'location', 'ubicacion_o_departamento'],
                'fecha_compra': ['fecha_compra', 'fecha de compra', 'purchase_date', 'fecha_compra'],
                'valor_compra': ['valor_compra', 'valor de compra', 'purchase_value', 'cost'],
                'empleado_responsable': ['empleado_responsable', 'responsable', 'assigned_to'],
                'observaciones': ['observaciones', 'observación', 'comments', 'notes'],
                'mac_address': ['mac_address', 'mac', 'direccion_mac'],
                'fecha_fabricacion': ['fecha_fabricacion', 'fecha_fabricación', 'manufacturing_date'],
                'vida_util_anios': ['vida_util_anios', 'vida_util_en_anios', 'useful_life'],
                'nombre_red': ['nombre_red', 'nombre_de_red', 'network_name'],
                'proveedor': ['proveedor', 'supplier'],
                'fecha_recepcion': ['fecha_recepcion', 'fecha_recepción', 'reception_date'],
                'fecha_mantenimiento': ['fecha_mantenimiento', 'maintenance_date'],
                'size_pantalla': ['size_pantalla', 'pantalla_tamaño', 'screen_size'],
                'resolucion': ['resolucion', 'resolución', 'resolution'],
                'procesador_marca': ['procesador_marca', 'processor_brand'],
                'procesador_velocidad': ['procesador_velocidad', 'processor_speed'],
                'procesador_generacion': ['procesador_generacion', 'processor_generation'],
                'sistema_operativo': ['sistema_operativo', 'operating_system'],
                'sistema_operativo_version': ['sistema_operativo_version', 'os_version'],
                'sistema_operativo_bits': ['sistema_operativo_bits', 'sistema_operativo_manejo_bit', 'os_bits'],
                'almacenamiento_capacidad': ['almacenamiento_capacidad', 'storage_capacity'],
                'memoria': ['memoria', 'memory'],
                'impresora_tipo': ['impresora_tipo', 'printer_type'],
                'impresora_velocidad_ppm': ['impresora_velocidad_ppm', 'printer_speed_ppm'],
                'impresora_color': ['impresora_color', 'printer_color'],
                'impresora_conexion': ['impresora_conexion', 'printer_connection'],
                'cpu_formato_diseno': ['cpu_formato_diseno', 'cpu_formato_diseño', 'cpu_design_format'],
                'proyector_lumens': ['proyector_lumens', 'projector_lumens'],
                'ups_vatios': ['ups_vatios', 'ups_watts'],
                'ups_fecha_bateria': ['ups_fecha_bateria', 'ups_battery_date'],
                'scanner_velocidad': ['scanner_velocidad', 'scanner_speed'],
                'scanner_color': ['scanner_color', 'scanner_color_bw'],
                'pantalla_proyector_tipo': ['pantalla_proyector_tipo', 'projector_screen_type'],
                'server_numero_procesadores': ['server_numero_procesadores', 'server_processor_count'],
                'licencia_tipo': ['licencia_tipo', 'license_type'],
                'licencia_clase': ['licencia_clase', 'license_class'],
                'mouse_tipo': ['mouse_tipo', 'mouse_type'],
                'mouse_conexion': ['mouse_conexion', 'mouse_connection'],
                'clase_disco': ['clase_disco', 'disk_class'],
               
            }

            # Obtener encabezados de la primera fila
            headers = [str(cell.value).lower().strip() if cell.value else '' for cell in ws[1]]
            if not headers:
                messages.error(request, 'El archivo Excel está vacío o no tiene encabezados.')
                return redirect('equipos')

            # Mapear encabezados a nombres de campos del modelo
            field_mapping = {}
            for model_field, possible_headers in header_map.items():
                for header in possible_headers:
                    if header.lower() in headers:
                        field_mapping[headers.index(header.lower())] = model_field
                        break

            if not field_mapping:
                messages.error(request, 'No se encontraron encabezados válidos en el archivo Excel.')
                return redirect('equipos')

            # Procesar cada fila
            for row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                if not any(row):  # Saltar filas vacías
                    continue

                try:
                    # Crear diccionario de datos basado en el mapeo
                    data = {}
                    for col_idx, model_field in field_mapping.items():
                        value = row[col_idx] if col_idx < len(row) else None

                        # Convertir fechas
                        if model_field == 'fecha_compra' and value:
                            if isinstance(value, str):
                                try:
                                    value = datetime.strptime(value, '%Y-%m-%d').date()
                                except ValueError:
                                    try:
                                        value = datetime.strptime(value, '%d/%m/%Y').date()
                                    except ValueError:
                                        errors.append(f'Fila {row_num}: Formato de fecha_compra inválido ({value}). Use YYYY-MM-DD o DD/MM/YYYY.')
                                        continue
                            elif isinstance(value, datetime):
                                value = value.date()
                            elif value is None:
                                value = None
                            else:
                                errors.append(f'Fila {row_num}: fecha_compra inválida ({value}).')
                                continue

                        # Convertir valores numéricos
                        if model_field == 'valor_compra' and value:
                            try:
                                value = float(value) if value is not None else None
                            except (ValueError, TypeError):
                                errors.append(f'Fila {row_num}: valor_compra debe ser un número ({value}).')
                                continue

                        data[model_field] = value

                    # Usar EquipoForm para validar
                    form = EquipoForm(data)
                    if form.is_valid():
                        equipo = form.save()
                        # Crear registro en HistorialEquipo
                        HistorialEquipo.objects.create(
                            equipo=equipo,
                            accion='creado',
                            usuario=request.user,
                            tipo=data.get('tipo', equipo.tipo),
                            marca=data.get('marca', equipo.marca),
                            modelo=data.get('modelo', equipo.modelo),
                            serial=data.get('serial', equipo.serial)
                        )
                        success_count += 1
                    else:
                        errors.append(f'Fila {row_num}: {form.errors.as_text()}')
                except Exception as e:
                    errors.append(f'Fila {row_num}: Error al procesar ({str(e)}).')

            if success_count > 0:
                messages.success(request, f'{success_count} equipos importados exitosamente.')
            if errors:
                messages.error(request, 'Errores en las siguientes filas: ' + '; '.join(errors))
                return redirect('equipos')

        except Exception as e:
            messages.error(request, f'Error al cargar el archivo Excel: {str(e)}')
            return redirect('equipos')

    return render(request, 'Equipos/importar_excel.html')
